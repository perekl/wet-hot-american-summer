#!/usr/bin/env python3
"""Generate WHAS table-read production kit from screenplay-derived cue data."""

from __future__ import annotations

import json
from collections import OrderedDict
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from asset_library import CANONICAL_ASSETS, resolve_asset_key, suggested_volume_for_mode
from background_cues import enrich_background_fields
from cues_data import CUES, LICENSED_MUSIC
from generate_annotated_screenplay import generate_annotated_screenplay

ROOT = Path(__file__).resolve().parent.parent
DATA = ROOT / "data"
DOCS = ROOT / "docs"

CUE_COLS = [
    "Cue ID", "Asset ID", "Cue Type", "Expected Background Asset",
    "Script Page", "Scene", "Trigger Dialogue", "Cue Name",
    "Category", "Priority", "Duration", "Loop", "Fade", "Volume", "Notes",
]
ASSET_COLS = [
    "Asset ID", "Name", "Category", "Playback Mode", "Suggested Volume",
    "Royalty Free?", "Source", "Filename", "Duration",
    "Used By Cue IDs", "Status", "Download URL", "Notes",
]


def assign_cue_assets(cues: list[dict]) -> list[dict]:
    """Attach canonical Asset ID to each cue; cue Volume/Fade/Loop preserve behavior."""
    enriched = []
    for cue in cues:
        key = resolve_asset_key(cue)
        if key not in CANONICAL_ASSETS:
            raise KeyError(f"No canonical asset for key '{key}' (cue {cue['Cue ID']})")
        asset_id = _asset_id_for_key(key)
        row = dict(cue)
        row["Asset ID"] = asset_id
        enriched.append(row)
    return enriched


def _asset_id_for_key(key: str) -> str:
    keys = sorted(CANONICAL_ASSETS.keys())
    return f"AST-{keys.index(key) + 1:03d}"


def _key_by_asset_id(asset_id: str) -> str:
    keys = sorted(CANONICAL_ASSETS.keys())
    idx = int(asset_id.split("-")[1]) - 1
    return keys[idx]


def build_assets(cues: list[dict]) -> list[dict]:
    """Build consolidated asset rows from canonical library + cue usage."""
    usage: OrderedDict[str, list[str]] = OrderedDict()
    for cue in cues:
        key = resolve_asset_key(cue)
        usage.setdefault(key, []).append(cue["Cue ID"])

    assets = []
    for key in sorted(CANONICAL_ASSETS.keys()):
        if key not in usage:
            continue
        spec = CANONICAL_ASSETS[key]
        mode = spec["playback_mode"]
        assets.append({
            "Asset ID": _asset_id_for_key(key),
            "Name": spec["name"],
            "Category": spec["category"],
            "Playback Mode": mode,
            "Suggested Volume": suggested_volume_for_mode(mode),
            "Royalty Free?": spec["royalty_free"],
            "Source": spec["source"],
            "Filename": spec["filename"],
            "Duration": spec["duration"],
            "Used By Cue IDs": ", ".join(usage[key]),
            "Status": "Needed",
            "Download URL": "",
            "Notes": spec["notes"],
        })
    return assets


def style_header(ws, cols):
    fill = PatternFill("solid", fgColor="1F4E79")
    font = Font(bold=True, color="FFFFFF", size=11)
    for i, col in enumerate(cols, 1):
        cell = ws.cell(row=1, column=i, value=col)
        cell.fill = fill
        cell.font = font
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = f"A1:{get_column_letter(len(cols))}1"


def autosize_columns(ws, cols, rows):
    for i, col in enumerate(cols, 1):
        max_len = len(col)
        for row in rows:
            val = str(row.get(col, ""))
            max_len = max(max_len, min(len(val), 80))
        ws.column_dimensions[get_column_letter(i)].width = max_len + 2


def write_xlsx(path: Path, sheet_name: str, cols: list[str], rows: list[dict]):
    wb = Workbook()
    ws = wb.active
    ws.title = sheet_name
    style_header(ws, cols)
    for r, row in enumerate(rows, 2):
        for c, col in enumerate(cols, 1):
            cell = ws.cell(row=r, column=c, value=row.get(col, ""))
            cell.alignment = Alignment(vertical="top", wrap_text=True)
    autosize_columns(ws, cols, rows)
    path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(path)


def write_cue_book_docx(path: Path, cues: list[dict]):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    title = doc.add_heading("Wet Hot American Summer", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Sound Cue Book — Table Read Production")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(14)
    doc.add_paragraph("Source: Wet Hot American Summer_v2.pdf (screenplay)")
    doc.add_paragraph(f"Total Cues: {len(cues)}")
    doc.add_paragraph("")

    for cue in cues:
        line = doc.add_paragraph()
        line.add_run(f"{cue['Cue ID']}  |  p.{cue['Script Page']}  |  ").bold = True
        line.add_run(f"{cue['Cue Name']}").bold = True
        line.add_run(f"  [{cue['Category']} / {cue['Priority']}]")
        detail = doc.add_paragraph(style="List Bullet")
        detail.add_run(f"Scene: {cue['Scene']}\n")
        detail.add_run(f"Trigger: {cue['Trigger Dialogue']}\n")
        detail.add_run(
            f"Duration: {cue['Duration']} | Loop: {cue['Loop']} | Fade: {cue['Fade']} | Vol: {cue['Volume']}"
        )
        if cue["Notes"]:
            detail.add_run(f"\nNotes: {cue['Notes']}")

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def write_stage_manager_docx(path: Path, cues: list[dict]):
    doc = Document()
    title = doc.add_heading("STAGE MANAGER SOUND CUE BOOK", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("WET HOT AMERICAN SUMMER — LIVE TABLE READ").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    for cue in cues:
        p = doc.add_paragraph()
        run = p.add_run(f"▶  {cue['Cue ID']}  —  PAGE {cue['Script Page']}\n")
        run.font.size = Pt(16)
        run.bold = True
        run2 = p.add_run(f"{cue['Cue Name'].upper()}\n")
        run2.font.size = Pt(20)
        run2.bold = True
        run3 = p.add_run(f"GO ON: {cue['Trigger Dialogue']}\n")
        run3.font.size = Pt(14)
        run4 = p.add_run(
            f"{cue['Category']} | {cue['Priority']} | {cue['Fade']} | Vol {cue['Volume']}"
        )
        run4.font.size = Pt(12)
        doc.add_paragraph("")

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def write_cue_book_pdf(path: Path, cues: list[dict]):
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = SimpleDocTemplate(str(path), pagesize=letter, topMargin=0.5 * inch, bottomMargin=0.5 * inch)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("Title", parent=styles["Heading1"], fontSize=16, spaceAfter=6, alignment=1)
    sub_style = ParagraphStyle("Sub", parent=styles["Normal"], fontSize=10, spaceAfter=12, alignment=1)
    row_style = ParagraphStyle("Row", parent=styles["Normal"], fontSize=8, leading=10)
    story = [
        Paragraph("Wet Hot American Summer — Sound Cue Book", title_style),
        Paragraph(f"Table Read Production | {len(cues)} cues | Source: screenplay PDF", sub_style),
    ]
    data = [["Cue", "Pg", "Name", "Category", "Trigger", "Fade", "Vol"]]
    for c in cues:
        data.append([
            c["Cue ID"], str(c["Script Page"]), c["Cue Name"][:40],
            c["Category"], c["Trigger Dialogue"][:55], c["Fade"], c["Volume"],
        ])
    table = Table(data, colWidths=[0.55 * inch, 0.35 * inch, 1.4 * inch, 0.7 * inch, 2.5 * inch, 0.55 * inch, 0.4 * inch])
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1F4E79")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 7),
        ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F2F2F2")]),
    ]))
    story.append(table)
    story.append(Spacer(1, 12))
    for c in cues:
        story.append(Paragraph(
            f"<b>{c['Cue ID']}</b> p.{c['Script Page']} — {c['Cue Name']} "
            f"[{c['Category']}] — {c['Trigger Dialogue']}",
            row_style,
        ))
    doc.build(story)


def write_json_exports(cues: list[dict], assets: list[dict]):
    asset_by_id = {a["Asset ID"]: a for a in assets}
    cue_export = []
    for i, c in enumerate(cues):
        asset = asset_by_id[c["Asset ID"]]
        entry = {
            "index": i,
            "id": c["Cue ID"],
            "asset_id": c["Asset ID"],
            "asset_filename": asset["Filename"],
            "playback_mode": asset["Playback Mode"],
            "suggested_volume": asset["Suggested Volume"],
            "page": c["Script Page"],
            "scene": c["Scene"],
            "trigger": c["Trigger Dialogue"],
            "name": c["Cue Name"],
            "category": c["Category"],
            "priority": c["Priority"],
            "duration": c["Duration"],
            "loop": c["Loop"],
            "fade": c["Fade"],
            "volume": int(c["Volume"]) if str(c["Volume"]).isdigit() else c["Volume"],
            "notes": c["Notes"],
        }
        if c.get("Cue Type"):
            entry["cue_type"] = c["Cue Type"]
        if c.get("Expected Background Asset"):
            entry["expected_background_asset_id"] = c["Expected Background Asset"]
        cue_export.append(entry)
    asset_export = []
    for a in assets:
        asset_export.append({
            "id": a["Asset ID"],
            "name": a["Name"],
            "category": a["Category"],
            "playback_mode": a["Playback Mode"],
            "suggested_volume": a["Suggested Volume"],
            "royalty_free": a["Royalty Free?"],
            "source": a["Source"],
            "filename": a["Filename"],
            "duration": a["Duration"],
            "used_by": [x.strip() for x in a["Used By Cue IDs"].split(",")],
            "status": a["Status"],
            "download_url": a["Download URL"],
            "notes": a["Notes"],
        })
    DATA.mkdir(parents=True, exist_ok=True)
    (DATA / "cues.json").write_text(json.dumps(cue_export, indent=2), encoding="utf-8")
    (DATA / "assets.json").write_text(json.dumps(asset_export, indent=2), encoding="utf-8")


def verify_page_coverage(cues: list[dict], total_pages: int = 90):
  pages = {c["Script Page"] for c in cues}
  missing = [p for p in range(1, total_pages + 1) if p not in pages]
  if missing:
      raise SystemExit(f"Missing screenplay page coverage: {missing}")
  print(f"Page coverage: {len(pages)}/{total_pages} pages, {len(cues)} cues")


def main():
    cues = assign_cue_assets(CUES)
    cues = enrich_background_fields(cues, _asset_id_for_key)
    assets = build_assets(cues)
    verify_page_coverage(cues)

    write_xlsx(DATA / "master_cues.xlsx", "Master Cues", CUE_COLS, cues)
    write_xlsx(DATA / "master_sound_assets.xlsx", "Sound Assets", ASSET_COLS, assets)
    write_xlsx(
        DATA / "licensed_music.xlsx", "Licensed Music",
        ["Song", "Artist", "Scene", "Trigger", "Fade", "Notes"],
        LICENSED_MUSIC,
    )

    write_cue_book_docx(DOCS / "CueBook.docx", cues)
    write_stage_manager_docx(DOCS / "StageManagerBook.docx", cues)
    write_cue_book_pdf(DOCS / "CueBook.pdf", cues)

    write_json_exports(cues, assets)

    annotated = generate_annotated_screenplay()
    print(f"Annotated screenplay: {annotated}")

    reuse = len(cues) / len(assets)
    print(f"Generated {len(cues)} cues, {len(assets)} reusable assets ({reuse:.1f}x reuse), "
          f"{len(LICENSED_MUSIC)} licensed tracks")
    print(f"Output: {DATA}, {DOCS}")


if __name__ == "__main__":
    main()
